"""
Revolutionary FastAPI application for HandyWriterz backend.
Production-ready with comprehensive error handling and resilience patterns.
"""

import asyncio
import json
import logging
import os
import time
import uuid
from contextlib import asynccontextmanager
from typing import Dict, Any, Optional, List

import redis.asyncio as redis
import uvicorn
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse, FileResponse
from fastapi.exceptions import RequestValidationError
from fastapi.staticfiles import StaticFiles
from langchain_core.messages import HumanMessage
from pydantic import BaseModel

# Import agent system for HandyWriterz workflow
from agent.handywriterz_state import HandyWriterzState, UserParams
from agent.handywriterz_graph import handywriterz_graph

# Import simple system through organized module structure
try:
    # Correctly import from the top-level src directory
    from src.agent.graph import graph as simple_graph
    from src.agent.state import OverallState as SimpleState
    SIMPLE_SYSTEM_AVAILABLE = True
    logger.info("✅ Simple Gemini system imported successfully")
except ImportError as e:
    SIMPLE_SYSTEM_AVAILABLE = False
    logger.warning(f"⚠️  Simple Gemini system not available: {e}")
    simple_graph = None
    SimpleState = None

# Import routing system
from agent.routing import SystemRouter, UnifiedProcessor
from db.database import (
    get_database, get_user_repository, get_conversation_repository,
    get_document_repository, db_manager
)

# Import admin routes
from routes.admin_models import router as admin_models_router
from db.models import User, Conversation, Document
from services.error_handler import (
    error_handler, get_error_handler, ErrorContext, ErrorCategory, ErrorSeverity,
    with_error_handling, with_circuit_breaker, with_retry
)
from services.security_service import (
    security_service, get_security_service, get_current_user, require_authorization,
    require_rate_limit, validate_input
)
from middleware.error_middleware import (
    error_middleware, global_exception_handler
)
from middleware.security_middleware import (
    security_middleware, csrf_middleware
)
from scripts import init_database

# Configure production logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('handywriterz.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Async Redis client for SSE
redis_client = redis.from_url(os.getenv("REDIS_URL", "redis://localhost:6379"), decode_responses=True)


# Initialize unified processor
unified_processor = UnifiedProcessor(SIMPLE_SYSTEM_AVAILABLE, True)


from db.models import Base
from sqlalchemy import create_engine

# Create database tables on startup, outside of the lifespan manager
# to ensure they are created even if the lifespan event doesn't fire
# in all execution contexts (like some test runners or script executions).
try:
    db_url = os.getenv("DATABASE_URL")
    if db_url:
        if db_url.startswith("postgres://"):
            db_url = db_url.replace("postgres://", "postgresql://", 1)
        engine = create_engine(db_url)
        Base.metadata.create_all(bind=engine)
        logger.info("✅ Database tables created successfully on startup.")
    else:
        logger.warning("⚠️ DATABASE_URL not set, skipping table creation on startup.")
except Exception as e:
    logger.error(f"❌ Failed to create database tables on startup: {e}")


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Manage application lifespan with comprehensive health checks."""
    logger.info("Starting HandyWriterz Revolutionary Backend...")

    # Initialize database with system prompts
    try:
        await init_database.main()
        logger.info("✅ Database initialized with system prompts.")
    except Exception as e:
        logger.error(f"❌ Failed to initialize database with system prompts: {e}")

    # Initialize systems with error handling
    startup_errors = []
    
    # Test Redis connection
    try:
        await redis_client.ping()
        logger.info("✅ Redis connection successful")
    except Exception as e:
        startup_errors.append(f"Redis connection failed: {e}")
        logger.error(f"❌ Redis connection failed: {e}")
    
    # Test Database connection
    try:
        if db_manager.health_check():
            logger.info("✅ Database connection successful")
        else:
            startup_errors.append("Database health check failed")
            logger.error("❌ Database health check failed")
    except Exception as e:
        startup_errors.append(f"Database connection failed: {e}")
        logger.error(f"❌ Database connection failed: {e}")
    
    # Test Error Handler
    try:
        await error_handler.get_error_statistics()
        logger.info("✅ Error handler initialized successfully")
    except Exception as e:
        startup_errors.append(f"Error handler initialization failed: {e}")
        logger.error(f"❌ Error handler initialization failed: {e}")
    
    # Log startup summary
    if startup_errors:
        logger.warning(f"Backend started with {len(startup_errors)} issues:")
        for error in startup_errors:
            logger.warning(f"  - {error}")
    else:
        logger.info("🚀 All systems operational - HandyWriterz backend ready!")
    
    yield
    
    logger.info("Shutting down HandyWriterz backend...")
    
    # Graceful shutdown
    shutdown_tasks = []
    
    # Close database connections
    try:
        db_manager.close()
        logger.info("✅ Database connections closed")
    except Exception as e:
        logger.error(f"❌ Error closing database: {e}")
    
    # Close Redis connections
    try:
        await redis_client.close()
        logger.info("✅ Redis connections closed")
    except Exception as e:
        logger.error(f"❌ Error closing Redis: {e}")
    
    logger.info("🔄 HandyWriterz backend shutdown complete")


# Create FastAPI app with enhanced configuration
app = FastAPI(
    title="Unified AI Platform - Revolutionary API",
    description="🚀 Intelligent Multi-Agent System: Simple Gemini + Advanced HandyWriterz with Automatic Routing",
    version="2.0.0-unified",
    lifespan=lifespan,
    docs_url="/docs",
    redoc_url="/redoc",
    openapi_url="/openapi.json"
)

# Add Security Middleware (first layer of protection)
app.add_middleware(security_middleware)

# Add CSRF Protection Middleware
app.add_middleware(csrf_middleware)

# Add Error Handling Middleware
app.add_middleware(error_middleware)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3000",
        "http://localhost:3001",
        "http://localhost:5173",  # SvelteKit dev server
        "https://handywriterz.vercel.app",  # Production frontend
        "https://*.handywriterz.com"  # Custom domains
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["X-Request-ID", "X-Response-Time", "X-Error-ID", "X-Security-Middleware"]
)

# Add global exception handlers
app.add_exception_handler(HTTPException, global_exception_handler.handle_http_exception)
app.add_exception_handler(RequestValidationError, global_exception_handler.handle_validation_exception)

# Include admin routes
app.include_router(admin_models_router)

# Include payment system routes
from api.payments import router as payments_router
from api.payout import router as payout_router
from api.checker import router as checker_router

app.include_router(payments_router)
app.include_router(payout_router)
app.include_router(checker_router)

# Mount static files for serving the SvelteKit frontend
static_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "static")
if os.path.exists(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")
    logger.info(f"✅ Static files mounted from {static_dir}")
else:
    logger.warning(f"⚠️ Static directory not found: {static_dir}")

# Mount pyodide files if they exist
pyodide_dir = os.path.join(static_dir, "pyodide") if os.path.exists(static_dir) else None
if pyodide_dir and os.path.exists(pyodide_dir):
    app.mount("/pyodide", StaticFiles(directory=pyodide_dir), name="pyodide")
    logger.info(f"✅ Pyodide files mounted from {pyodide_dir}")

# Mount SvelteKit build directory for frontend serving
build_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "build")
if os.path.exists(build_dir):
    app.mount("/app", StaticFiles(directory=build_dir), name="frontend_build")
    logger.info(f"✅ SvelteKit build mounted from {build_dir}")
else:
    logger.warning(f"⚠️ SvelteKit build directory not found: {build_dir}")


# Pydantic models
class WritingRequest(BaseModel):
    """Request model for academic writing."""
    prompt: str
    user_params: Dict[str, Any]
    auth_token: Optional[str] = None
    payment_transaction_id: Optional[str] = None
    uploaded_file_urls: List[str] = []


class WritingResponse(BaseModel):
    """Response model for writing request."""
    conversation_id: str
    status: str
    message: str


class HealthResponse(BaseModel):
    """Health check response."""
    status: str
    timestamp: float
    version: str


# Health check endpoints
@app.get("/health", response_model=HealthResponse)
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.MEDIUM)
async def health_check():
    """Comprehensive health check endpoint."""
    return HealthResponse(
        status="healthy",
        timestamp=time.time(),
        version="2.0.0-unified"
    )


# Enhanced system status endpoint for unified platform
@app.get("/api/status")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def unified_system_status():
    """
    🎯 Unified System Status - Enhanced with routing information
    Shows the status of both simple and advanced systems with routing capabilities.
    """
    try:
        # Get routing statistics
        routing_stats = unified_processor.router.get_routing_stats()
        
        # Test system availability
        simple_status = "available" if SIMPLE_SYSTEM_AVAILABLE else "unavailable"
        advanced_status = "available"  # HandyWriterz is always available in this context
        
        # Get Redis status
        redis_status = "unknown"
        try:
            await redis_client.ping()
            redis_status = "healthy"
        except:
            redis_status = "unhealthy"
        
        # Get database status
        db_status = "unknown"
        try:
            if db_manager.health_check():
                db_status = "healthy"
            else:
                db_status = "unhealthy"
        except:
            db_status = "unavailable"
        
        return {
            "status": "operational",
            "version": "2.0.0-unified",
            "timestamp": time.time(),
            "platform": "Unified AI Platform",
            
            # System availability
            "systems": {
                "simple_gemini": {
                    "status": simple_status,
                    "description": "Quick responses for simple queries",
                    "capabilities": ["chat", "basic_research", "quick_answers"]
                },
                "advanced_handywriterz": {
                    "status": advanced_status,
                    "description": "Full academic writing workflow with 30+ agents",
                    "capabilities": [
                        "academic_writing", "research_swarms", "qa_swarms", 
                        "writing_swarms", "citation_management", "turnitin_checking",
                        "learning_outcomes", "multi_agent_orchestration"
                    ]
                }
            },
            
            # Intelligent routing
            "routing": {
                "enabled": True,
                "algorithm": "complexity_analysis",
                "thresholds": routing_stats["thresholds"],
                "modes": routing_stats["routing_modes"],
                "capabilities": routing_stats["capabilities"]
            },
            
            # Infrastructure
            "infrastructure": {
                "redis": {
                    "status": redis_status,
                    "purpose": "Caching and SSE"
                },
                "database": {
                    "status": db_status,
                    "purpose": "User data and conversations"
                },
                "vector_storage": {
                    "status": "available",
                    "purpose": "Semantic search and embeddings"
                }
            },
            
            # Features
            "features": {
                "intelligent_routing": True,
                "multimodal_processing": True,
                "file_upload": True,
                "streaming_responses": True,
                "swarm_intelligence": advanced_status == "available",
                "academic_writing": advanced_status == "available",
                "web3_authentication": True,
                "real_time_collaboration": True,
                "citation_management": True,
                "plagiarism_checking": True
            },
            
            # API endpoints
            "endpoints": {
                "unified_chat": "/api/chat",
                "simple_chat": "/api/chat/simple",
                "advanced_chat": "/api/chat/advanced", 
                "academic_writing": "/api/write",
                "file_upload": "/api/upload",
                "streaming": "/api/stream/{conversation_id}",
                "documentation": "/docs"
            },
            
            # Performance metrics
            "performance": {
                "routing_overhead": "< 50ms",
                "simple_response_time": "1-3 seconds",
                "advanced_response_time": "30-300 seconds",
                "hybrid_response_time": "30-300 seconds (parallel processing)"
            }
        }
        
    except Exception as e:
        logger.error(f"System status check failed: {e}")
        return {
            "status": "degraded",
            "version": "2.0.0-unified",
            "timestamp": time.time(),
            "error": str(e),
            "systems": {
                "simple_gemini": {"status": "unknown"},
                "advanced_handywriterz": {"status": "unknown"}
            }
        }


# Routing analysis endpoint for development/debugging
@app.post("/api/analyze")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def analyze_request_complexity(
    message: str = Form(...),
    files: Optional[List[UploadFile]] = File(None),
    user_params: Optional[str] = Form(None)
):
    """
    🔍 Analyze Request Complexity - Development endpoint
    Shows how the routing system would handle a request without actually processing it.
    """
    try:
        # Process files for analysis
        processed_files = []
        if files:
            for file in files:
                content = await file.read()
                processed_files.append({
                    "filename": file.filename,
                    "size": len(content),
                    "type": file.content_type
                })
        
        # Parse user parameters
        parsed_user_params = {}
        if user_params:
            try:
                parsed_user_params = json.loads(user_params)
            except json.JSONDecodeError:
                pass
        
        # Analyze routing
        routing = await unified_processor.router.analyze_request(
            message, processed_files, parsed_user_params
        )
        
        # Get detailed analysis
        return {
            "message": message,
            "analysis": {
                "word_count": len(message.split()),
                "file_count": len(processed_files),
                "has_user_params": bool(parsed_user_params),
                "academic_indicators": [
                    keyword for keyword in [
                        "essay", "research", "academic", "citation", "thesis", 
                        "dissertation", "literature review", "methodology"
                    ] if keyword.lower() in message.lower()
                ]
            },
            "routing_decision": routing,
            "explanation": {
                "simple": "Fast responses for quick questions and simple tasks",
                "advanced": "Full academic workflow with research, writing, and quality assurance",
                "hybrid": "Parallel processing with both systems for comprehensive results"
            },
            "estimated_processing_time": {
                "simple": "1-3 seconds",
                "advanced": "30-300 seconds", 
                "hybrid": "30-300 seconds (parallel)"
            }
        }
        
    except Exception as e:
        logger.error(f"Request analysis failed: {e}")
        return {
            "error": str(e),
            "message": message,
            "routing_decision": {
                "system": "advanced",
                "complexity": 5.0,
                "reason": "analysis_failed",
                "confidence": 0.0
            }
        }


@app.get("/health/detailed")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.MEDIUM)
async def detailed_health_check():
    """Detailed health check with system status."""
    health_data = {
        "status": "healthy",
        "timestamp": time.time(),
        "version": "2.0.0",
        "uptime_seconds": (time.time() - app.state.start_time) if hasattr(app.state, 'start_time') else 0,
        "services": {}
    }
    
    # Check Redis
    try:
        await redis_client.ping()
        health_data["services"]["redis"] = {"status": "healthy", "response_time": "< 1ms"}
    except Exception as e:
        health_data["services"]["redis"] = {"status": "unhealthy", "error": str(e)}
        health_data["status"] = "degraded"
    
    # Check Database
    try:
        if db_manager.health_check():
            health_data["services"]["database"] = {"status": "healthy", "response_time": "< 10ms"}
        else:
            health_data["services"]["database"] = {"status": "unhealthy", "error": "Health check failed"}
            health_data["status"] = "degraded"
    except Exception as e:
        health_data["services"]["database"] = {"status": "unhealthy", "error": str(e)}
        health_data["status"] = "degraded"
    
    # Check Error Handler
    try:
        error_stats = await error_handler.get_error_statistics()
        health_data["services"]["error_handler"] = {
            "status": "healthy",
            "total_errors": error_stats["total_errors"],
            "critical_errors": error_stats["critical_errors"]
        }
    except Exception as e:
        health_data["services"]["error_handler"] = {"status": "unhealthy", "error": str(e)}
    
    return health_data


@app.get("/metrics")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_system_metrics():
    """Get comprehensive system metrics."""
    try:
        # Get error handler statistics
        error_stats = await error_handler.get_error_statistics()
        
        # Get middleware statistics
        middleware_instance = None
        for middleware in app.user_middleware:
            if hasattr(middleware, 'cls') and middleware.cls.__name__ == 'RevolutionaryErrorMiddleware':
                middleware_instance = middleware
                break
        
        middleware_stats = {}
        if middleware_instance and hasattr(middleware_instance, 'get_middleware_stats'):
            middleware_stats = await middleware_instance.get_middleware_stats()
        
        return {
            "timestamp": time.time(),
            "system": {
                "version": "2.0.0",
                "uptime": error_stats.get("uptime_seconds", 0)
            },
            "errors": error_stats,
            "middleware": middleware_stats,
            "services": {
                "redis": {"status": "healthy" if await _check_redis() else "unhealthy"},
                "database": {"status": "healthy" if db_manager.health_check() else "unhealthy"}
            }
        }
    except Exception as e:
        logger.error(f"Failed to get system metrics: {e}")
        raise HTTPException(status_code=500, detail="Failed to retrieve system metrics")


async def _check_redis() -> bool:
    """Check Redis health."""
    try:
        await redis_client.ping()
        return True
    except:
        return False


# Authentication and Security endpoints
@app.post("/api/auth/login")
@require_rate_limit("authentication")
@validate_input()
@with_error_handling(ErrorCategory.AUTHENTICATION, ErrorSeverity.MEDIUM)
async def login(
    request: Request,
    login_data: Dict[str, Any],
    user_repo=Depends(get_user_repository)
):
    """Authenticate user and return JWT token."""
    wallet_address = login_data.get("wallet_address")
    
    if not wallet_address:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Wallet address is required"
        )
    
    # Get or create user
    user = user_repo.get_user_by_wallet(wallet_address)
    if not user:
        # Create new user for first-time login
        user = user_repo.create_user(
            wallet_address=wallet_address,
            user_type="student",
            subscription_tier="free"
        )
    
    # Generate JWT token
    token = await security_service.generate_jwt_token(user.to_dict())
    
    logger.info(f"User authenticated: {wallet_address}")
    
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": str(user.id),
            "wallet_address": user.wallet_address,
            "user_type": user.user_type.value if user.user_type else "student",
            "subscription_tier": user.subscription_tier,
            "credits_remaining": user.credits_remaining
        }
    }


@app.post("/api/auth/verify")
@with_error_handling(ErrorCategory.AUTHENTICATION, ErrorSeverity.LOW)
async def verify_token(current_user: Dict[str, Any] = Depends(get_current_user)):
    """Verify JWT token and return user information."""
    return {
        "valid": True,
        "user": current_user
    }


# Unified Chat Endpoint with Intelligent Routing
@app.post("/api/chat")
@require_rate_limit("chat_request")
@validate_input()
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.MEDIUM)
async def unified_chat_endpoint(
    message: str = Form(...),
    files: Optional[List[UploadFile]] = File(None),
    user_params: Optional[str] = Form(None),  # JSON string of user parameters
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user)
):
    """
    🚀 Unified Chat Endpoint with Intelligent Routing
    
    Automatically routes between:
    - Simple Gemini system (quick responses) 
    - Advanced HandyWriterz system (academic writing)
    - Hybrid mode (both systems in parallel)
    
    Routing is based on request complexity analysis.
    """
    try:
        # Process uploaded files
        processed_files = []
        if files:
            for file in files:
                content = await file.read()
                processed_files.append({
                    "filename": file.filename,
                    "content_type": file.content_type,
                    "size": len(content),
                    "content": content.decode('utf-8', errors='ignore') if file.content_type and file.content_type.startswith("text") else content
                })
        
        # Parse user parameters if provided
        parsed_user_params = {}
        if user_params:
            try:
                parsed_user_params = json.loads(user_params)
            except json.JSONDecodeError:
                logger.warning(f"Invalid user_params JSON: {user_params}")
        
        # Get user ID if available
        user_id = str(current_user.get("id")) if current_user else None
        
        # Process using unified processor
        result = await unified_processor.process_message(
            message=message,
            files=processed_files,
            user_params=parsed_user_params,
            user_id=user_id
        )
        
        # Enhanced response format compatible with both frontend expectations
        response = {
            "success": result.get("success", True),
            "response": result.get("response", ""),
            "sources": result.get("sources", []),
            "workflow_status": result.get("workflow_status", "completed"),
            
            # Routing information
            "system_used": result.get("system_used", "unknown"),
            "complexity_score": result.get("complexity_score", 0.0),
            "routing_reason": result.get("routing_reason", ""),
            "routing_confidence": result.get("routing_confidence", 0.0),
            "processing_time": result.get("processing_time", 0.0),
            
            # Advanced features (when available)
            "conversation_id": result.get("conversation_id"),
            "quality_score": result.get("quality_score"),
            "agent_metrics": result.get("agent_metrics", {}),
            "citation_count": result.get("citation_count", 0),
            "system_type": result.get("system_type", ""),
            
            # Hybrid mode specific
            "simple_insights": result.get("simple_insights"),
            "research_depth": result.get("research_depth", 0),
            "hybrid_results": result.get("hybrid_results", {})
        }
        
        # Log successful routing
        logger.info(f"✅ Chat processed successfully:")
        logger.info(f"   System: {result.get('system_used')}")
        logger.info(f"   Complexity: {result.get('complexity_score'):.1f}")
        logger.info(f"   Time: {result.get('processing_time'):.2f}s")
        
        return response
        
    except Exception as e:
        logger.error(f"Unified chat endpoint error: {e}")
        
        # Enhanced error response
        error_response = {
            "success": False,
            "response": f"I encountered an error processing your request: {str(e)}",
            "sources": [],
            "workflow_status": "failed",
            "system_used": "error_fallback",
            "complexity_score": 0.0,
            "error_details": {
                "error_type": type(e).__name__,
                "error_message": str(e)
            }
        }
        
        # Return error response instead of raising HTTP exception for better UX
        return error_response


# Quick chat endpoint for simple queries (explicit routing)
@app.post("/api/chat/simple")
@require_rate_limit("chat_request")
@validate_input()
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.MEDIUM)
async def simple_chat_endpoint(
    message: str = Form(...),
    files: Optional[List[UploadFile]] = File(None)
):
    """
    Simple chat endpoint - forces routing to Gemini system for quick responses.
    Use this for fast, lightweight queries.
    """
    if not SIMPLE_SYSTEM_AVAILABLE:
        raise HTTPException(
            status_code=503,
            detail="Simple system not available. Use /api/chat for automatic routing."
        )
    
    try:
        # Process files
        processed_files = []
        if files:
            for file in files:
                content = await file.read()
                processed_files.append({
                    "filename": file.filename,
                    "content": content.decode('utf-8', errors='ignore') if file.content_type and file.content_type.startswith("text") else content
                })
        
        # Force simple processing
        result = await unified_processor._process_simple(message, processed_files)
        
        return {
            "success": True,
            "response": result.get("response", ""),
            "sources": result.get("sources", []),
            "system_used": "simple_forced",
            "processing_time": 0.0  # Would need to time this
        }
        
    except Exception as e:
        logger.error(f"Simple chat endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Advanced chat endpoint for academic writing (explicit routing)  
@app.post("/api/chat/advanced")
@require_rate_limit("create_document")
@validate_input()
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.HIGH)
async def advanced_chat_endpoint(
    message: str = Form(...),
    files: Optional[List[UploadFile]] = File(None),
    user_params: Optional[str] = Form(None),
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user)
):
    """
    Advanced chat endpoint - forces routing to HandyWriterz system for academic writing.
    Use this for complex academic writing tasks.
    """
    try:
        # Process files
        processed_files = []
        if files:
            for file in files:
                content = await file.read()
                processed_files.append({
                    "filename": file.filename,
                    "content": content.decode('utf-8', errors='ignore') if file.content_type and file.content_type.startswith("text") else content
                })
        
        # Parse user parameters
        parsed_user_params = {}
        if user_params:
            try:
                parsed_user_params = json.loads(user_params)
            except json.JSONDecodeError:
                logger.warning(f"Invalid user_params JSON: {user_params}")
        
        user_id = str(current_user.get("id")) if current_user else None
        
        # Force advanced processing
        result = await unified_processor._process_advanced(
            message, processed_files, parsed_user_params, user_id
        )
        
        return {
            "success": True,
            "response": result.get("response", ""),
            "conversation_id": result.get("conversation_id"),
            "sources": result.get("sources", []),
            "quality_score": result.get("quality_score", 0),
            "system_used": "advanced_forced",
            "processing_time": 0.0  # Would need to time this
        }
        
    except Exception as e:
        logger.error(f"Advanced chat endpoint error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/security/stats")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_security_stats(
    current_user: Dict[str, Any] = Depends(require_authorization("admin_access"))
):
    """Get security statistics (admin only)."""
    try:
        # Get security events from Redis
        security_events = await redis_client.lrange("security_events", 0, 99)
        
        # Get middleware stats
        middleware_stats = {}
        for middleware in app.user_middleware:
            if hasattr(middleware, 'cls') and hasattr(middleware.cls, 'get_security_stats'):
                instance = getattr(middleware, 'instance', None)
                if instance and hasattr(instance, 'get_security_stats'):
                    middleware_stats = await instance.get_security_stats()
                    break
        
        return {
            "security_events": [json.loads(event) for event in security_events],
            "middleware_stats": middleware_stats,
            "total_events": len(security_events),
            "timestamp": time.time()
        }
        
    except Exception as e:
        logger.error(f"Failed to get security stats: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve security statistics"
        )


# Main writing endpoint
@app.post("/api/write", response_model=WritingResponse)
@require_rate_limit("create_document")
@validate_input()
@with_circuit_breaker("writing_workflow")
@with_error_handling(ErrorCategory.AGENT_FAILURE, ErrorSeverity.HIGH)
async def start_writing(
    request: WritingRequest,
    http_request: Request,
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user),
    user_repo=Depends(get_user_repository),
    conversation_repo=Depends(get_conversation_repository)
):
    """Start the academic writing process with comprehensive error handling."""
    context = ErrorContext(
        request_id=getattr(http_request.state, 'request_id', str(uuid.uuid4())),
        additional_data={
            "prompt_length": len(request.prompt),
            "user_params": request.user_params,
            "file_count": len(request.uploaded_file_urls)
        }
    )
    
    try:
        # Create or get user from wallet address
        user = None
        if request.auth_token:  # If user is authenticated
            # Extract wallet address from auth token (simplified)
            wallet_address = request.auth_token  # Placeholder - implement proper JWT parsing
            user = user_repo.get_user_by_wallet(wallet_address)
            if not user:
                user = user_repo.create_user(
                    wallet_address=wallet_address,
                    user_type="student",
                    subscription_tier="free"
                )
        
        # Validate user parameters
        try:
            user_params = UserParams(**request.user_params)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid user parameters: {e}")
        
        # Create conversation in database
        conversation = conversation_repo.create_conversation(
            user_id=str(user.id) if user else None,
            user_params=user_params.dict(),
            title=f"{user_params.writeupType.title()} - {user_params.field.title()}",
            workflow_status="initiated"
        )
        
        conversation_id = str(conversation.id)
        
        # Create initial state
        initial_state = HandyWriterzState(
            conversation_id=conversation_id,
            user_id=str(user.id) if user else "",
            wallet_address=user.wallet_address if user else None,
            messages=[HumanMessage(content=request.prompt)],
            user_params=user_params.dict(),
            uploaded_docs=[],
            outline=None,
            research_agenda=[],
            search_queries=[],
            raw_search_results=[],
            filtered_sources=[],
            verified_sources=[],
            draft_content=None,
            current_draft=None,
            revision_count=0,
            evaluation_results=[],
            evaluation_score=None,
            turnitin_reports=[],
            turnitin_passed=False,
            formatted_document=None,
            learning_outcomes_report=None,
            download_urls={},
            current_node=None,
            workflow_status="initiated",
            error_message=None,
            retry_count=0,
            max_iterations=5,
            enable_tutor_review=False,
            start_time=time.time(),
            end_time=None,
            processing_metrics={},
            auth_token=request.auth_token,
            payment_transaction_id=request.payment_transaction_id,
            uploaded_files=[{"url": url} for url in request.uploaded_file_urls]
        )
        
        # Start the workflow asynchronously
        asyncio.create_task(execute_writing_workflow(conversation_id, initial_state))
        
        return WritingResponse(
            conversation_id=conversation_id,
            status="started",
            message="Revolutionary academic writing process initiated. Connect to the stream endpoint for real-time updates."
        )
        
    except Exception as e:
        logger.error(f"Failed to start writing process: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# SSE streaming endpoint
@app.get("/api/stream/{conversation_id}")
async def stream_updates(conversation_id: str):
    """Stream real-time updates for a conversation."""
    
    async def generate_events():
        """Generate SSE events from Redis pub/sub."""
        pubsub = redis_client.pubsub()
        channel = f"sse:{conversation_id}"
        
        try:
            await pubsub.subscribe(channel)
            logger.info(f"Subscribed to SSE channel: {channel}")
            
            # Send initial connection event
            yield f"data: {json.dumps({'type': 'connected', 'conversation_id': conversation_id})}\n\n"
            
            # Listen for messages
            async for message in pubsub.listen():
                if message["type"] == "message":
                    try:
                        # Parse the message data safely - SECURITY FIX
                        event_data = json.loads(message["data"])
                        yield f"data: {json.dumps(event_data)}\n\n"
                        
                        # Break if workflow is complete or failed
                        if event_data.get("type") in ["workflow_complete", "workflow_failed"]:
                            break
                            
                    except Exception as e:
                        logger.error(f"Error processing SSE message: {e}")
                        continue
                        
        except Exception as e:
            logger.error(f"SSE stream error: {e}")
            yield f"data: {json.dumps({'type': 'error', 'message': str(e)})}\n\n"
        
        finally:
            await pubsub.unsubscribe(channel)
            logger.info(f"Unsubscribed from SSE channel: {channel}")
    
    return StreamingResponse(
        generate_events(),
        media_type="text/plain",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "Content-Type": "text/event-stream",
        }
    )



# File upload endpoint
@app.post("/api/upload")
@require_rate_limit("upload_file")
@with_error_handling(ErrorCategory.VALIDATION, ErrorSeverity.MEDIUM)
async def upload_file(
    file: UploadFile = File(...),
    current_user: Optional[Dict[str, Any]] = Depends(get_current_user)
):
    """Upload a file for processing with cloud storage."""
    try:
        # Validate file type and size
        allowed_types = [".pdf", ".docx", ".txt", ".md"]
        file_extension = os.path.splitext(file.filename)[1].lower()
        max_size = 10 * 1024 * 1024  # 10MB limit
        
        if file_extension not in allowed_types:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {file_extension}. Allowed types: {allowed_types}"
            )
        
        # Read file content
        content = await file.read()
        
        if len(content) > max_size:
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size: {max_size // (1024*1024)}MB"
            )
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        safe_filename = f"{file_id}_{file.filename}"
        
        # Upload to cloud storage (R2/S3)
        try:
            import boto3
            from botocore.config import Config
            
            # Configure R2 client (Cloudflare R2 is S3-compatible)
            r2_client = boto3.client(
                's3',
                endpoint_url=os.getenv('R2_ENDPOINT_URL'),
                aws_access_key_id=os.getenv('R2_ACCESS_KEY_ID'),
                aws_secret_access_key=os.getenv('R2_SECRET_ACCESS_KEY'),
                config=Config(signature_version='s3v4'),
                region_name='auto'
            )
            
            bucket_name = os.getenv('R2_BUCKET_NAME', 'handywriterz-uploads')
            
            # Upload file
            r2_client.put_object(
                Bucket=bucket_name,
                Key=safe_filename,
                Body=content,
                ContentType=file.content_type or 'application/octet-stream',
                Metadata={
                    'original_filename': file.filename,
                    'file_extension': file_extension,
                    'upload_timestamp': str(time.time())
                }
            )
            
            # Generate public URL
            file_url = f"https://{bucket_name}.r2.dev/{safe_filename}"
            
        except Exception as upload_error:
            logger.warning(f"Cloud upload failed, using local fallback: {upload_error}")
            # Fallback to local storage for development
            upload_dir = os.getenv('UPLOAD_DIR', '/tmp/uploads')
            os.makedirs(upload_dir, exist_ok=True)
            
            file_path = os.path.join(upload_dir, safe_filename)
            with open(file_path, 'wb') as f:
                f.write(content)
            
            file_url = f"/uploads/{safe_filename}"
        
        # Extract text content for processing
        extracted_text = ""
        try:
            if file_extension == ".txt":
                extracted_text = content.decode('utf-8')
            elif file_extension == ".md":
                extracted_text = content.decode('utf-8')
            elif file_extension == ".pdf":
                import PyPDF2
                from io import BytesIO
                pdf_reader = PyPDF2.PdfReader(BytesIO(content))
                extracted_text = "\n".join(page.extract_text() for page in pdf_reader.pages)
            elif file_extension == ".docx":
                import mammoth
                from io import BytesIO
                result = mammoth.convert_to_html(BytesIO(content))
                # For simplicity, we'll just use the text content.
                # A more advanced implementation could handle the HTML.
                text_result = mammoth.extract_raw_text(BytesIO(content))
                extracted_text = text_result.value
        except Exception as e:
            logger.warning(f"Text extraction failed: {e}")
        
        return {
            "file_id": file_id,
            "filename": file.filename,
            "size": len(content),
            "type": file.content_type,
            "url": file_url,
            "extracted_text": extracted_text[:1000] if extracted_text else "",  # Preview
            "status": "uploaded"
        }
        
    except Exception as e:
        logger.error(f"File upload failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Get conversation status
@app.get("/api/conversation/{conversation_id}")
async def get_conversation_status(
    conversation_id: str,
    conversation_repo=Depends(get_conversation_repository)
):
    """Get the current status of a conversation."""
    try:
        # Query database for real conversation status
        conversation = conversation_repo.get_conversation(conversation_id)
        
        if not conversation:
            raise HTTPException(
                status_code=404,
                detail=f"Conversation {conversation_id} not found"
            )
        
        # Calculate progress based on workflow status
        progress_map = {
            "initiated": 5.0,
            "planning": 15.0,
            "searching": 35.0,
            "filtering": 50.0,
            "writing": 70.0,
            "evaluating": 85.0,
            "formatting": 95.0,
            "completed": 100.0,
            "failed": 0.0
        }
        
        current_progress = progress_map.get(conversation.workflow_status, 0.0)
        
        # Estimate completion time based on status
        time_estimates = {
            "initiated": "8-12 minutes",
            "planning": "6-10 minutes", 
            "searching": "4-8 minutes",
            "filtering": "3-5 minutes",
            "writing": "2-4 minutes",
            "evaluating": "1-2 minutes",
            "formatting": "30 seconds",
            "completed": "Complete",
            "failed": "Failed"
        }
        
        return {
            "conversation_id": conversation_id,
            "status": conversation.workflow_status,
            "current_node": conversation.current_node,
            "progress": current_progress,
            "estimated_completion": time_estimates.get(conversation.workflow_status, "Unknown"),
            "created_at": conversation.created_at.isoformat(),
            "updated_at": conversation.updated_at.isoformat(),
            "user_params": conversation.user_params,
            "error_message": conversation.error_message,
            "retry_count": conversation.retry_count or 0
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get conversation status: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Download document endpoint
@app.get("/api/download/{conversation_id}/{document_type}")
async def download_document(
    conversation_id: str, 
    document_type: str,
    document_repo=Depends(get_document_repository)
):
    """Download a generated document."""
    try:
        # Validate document type
        allowed_types = ["docx", "txt", "pdf", "lo_report"]
        if document_type not in allowed_types:
            raise HTTPException(
                status_code=400, 
                detail=f"Invalid document type: {document_type}. Allowed: {allowed_types}"
            )
        
        # Get document from database
        document = document_repo.get_by_conversation_and_type(conversation_id, document_type)
        
        if not document:
            raise HTTPException(
                status_code=404, 
                detail=f"Document of type {document_type} not found for conversation {conversation_id}"
            )
        
        # Check if document has a file URL (cloud storage)
        if document.file_urls and document_type in document.file_urls:
            file_url = document.file_urls[document_type]
            
            # For cloud storage, redirect to the file URL
            from fastapi.responses import RedirectResponse
            return RedirectResponse(url=file_url, status_code=302)
        
        # If no file URL, generate document content on-the-fly
        content = document.content
        if not content:
            raise HTTPException(
                status_code=404,
                detail="Document content not available"
            )
        
        # Generate appropriate response based on document type
        if document_type == "docx":
            from io import BytesIO
            from docx import Document as DocxDocument
            
            # Create DOCX from content
            doc = DocxDocument()
            for paragraph in content.split('\n\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)
            
            from fastapi.responses import StreamingResponse
            return StreamingResponse(
                BytesIO(bio.getvalue()),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={conversation_id}_{document_type}.docx"}
            )
            
        elif document_type == "txt":
            from fastapi.responses import Response
            return Response(
                content=content,
                media_type="text/plain",
                headers={"Content-Disposition": f"attachment; filename={conversation_id}_{document_type}.txt"}
            )
            
        elif document_type == "lo_report":
            # Learning outcomes report as JSON
            lo_data = document.metadata.get("learning_outcomes", {}) if document.metadata else {}
            import json
            
            from fastapi.responses import Response
            return Response(
                content=json.dumps(lo_data, indent=2),
                media_type="application/json",
                headers={"Content-Disposition": f"attachment; filename={conversation_id}_learning_outcomes.json"}
            )
        
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported document type: {document_type}")
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Document download failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# User management endpoints
@app.get("/api/users/{wallet_address}")
async def get_user_profile(
    wallet_address: str,
    user_repo=Depends(get_user_repository)
):
    """Get user profile by wallet address."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)
        
        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )
        
        return {
            "id": str(user.id),
            "wallet_address": user.wallet_address,
            "user_type": user.user_type,
            "subscription_tier": user.subscription_tier,
            "credits_balance": user.credits_balance,
            "credits_used": user.credits_used,
            "documents_created": user.documents_created,
            "avg_quality_score": user.avg_quality_score,
            "created_at": user.created_at.isoformat(),
            "last_active": user.last_active.isoformat() if user.last_active else None,
            "preferences": user.preferences
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get user profile: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.put("/api/users/{wallet_address}")
async def update_user_profile(
    wallet_address: str,
    user_data: Dict[str, Any],
    user_repo=Depends(get_user_repository)
):
    """Update user profile."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)
        
        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )
        
        # Update allowed fields
        allowed_fields = ['user_type', 'subscription_tier', 'preferences']
        update_data = {k: v for k, v in user_data.items() if k in allowed_fields}
        
        user_repo.update_user_stats(str(user.id), **update_data)
        
        return {"status": "updated", "message": "User profile updated successfully"}
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to update user profile: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/users/{wallet_address}/conversations")
async def get_user_conversations(
    wallet_address: str,
    limit: int = 20,
    user_repo=Depends(get_user_repository),
    conversation_repo=Depends(get_conversation_repository)
):
    """Get user's conversation history."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)
        
        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )
        
        conversations = conversation_repo.get_user_conversations(str(user.id), limit)
        
        return {
            "conversations": [
                {
                    "id": str(conv.id),
                    "title": conv.title,
                    "workflow_status": conv.workflow_status,
                    "created_at": conv.created_at.isoformat(),
                    "updated_at": conv.updated_at.isoformat(),
                    "user_params": conv.user_params,
                    "error_message": conv.error_message
                }
                for conv in conversations
            ]
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to get user conversations: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Credits management endpoints  
@app.get("/api/credits/{wallet_address}")
async def get_user_credits(
    wallet_address: str,
    user_repo=Depends(get_user_repository)
):
    """Get user's credit balance and usage statistics."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)
        
        if not user:
            # Create user if they don't exist
            user = user_repo.create_user(
                wallet_address=wallet_address,
                user_type="student",
                subscription_tier="free",
                credits_balance=3  # Welcome bonus
            )
        
        return {
            "wallet_address": wallet_address,
            "credits_balance": user.credits_balance,
            "credits_used": user.credits_used,
            "subscription_tier": user.subscription_tier,
            "next_renewal": None,  # TODO: Add subscription logic
            "usage_stats": {
                "documents_created": user.documents_created,
                "avg_quality_score": user.avg_quality_score,
                "total_words_generated": user.total_words_generated or 0
            }
        }
        
    except Exception as e:
        logger.error(f"Failed to get user credits: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/credits/{wallet_address}/purchase")
async def purchase_credits(
    wallet_address: str,
    purchase_data: Dict[str, Any],
    user_repo=Depends(get_user_repository)
):
    """Purchase credits for a user."""
    try:
        amount = purchase_data.get("amount", 0)
        payment_method = purchase_data.get("payment_method", "crypto")
        transaction_id = purchase_data.get("transaction_id")
        
        if amount <= 0:
            raise HTTPException(
                status_code=400,
                detail="Invalid credit amount"
            )
        
        user = user_repo.get_user_by_wallet(wallet_address)
        
        if not user:
            raise HTTPException(
                status_code=404,
                detail="User not found"
            )
        
        # TODO: Verify payment transaction
        # For now, we'll trust the transaction
        
        # Add credits to user balance
        new_balance = user.credits_balance + amount
        user_repo.update_user_stats(
            str(user.id),
            credits_balance=new_balance
        )
        
        logger.info(f"Credits purchased: {wallet_address} - {amount} credits")
        
        return {
            "status": "success",
            "credits_added": amount,
            "new_balance": new_balance,
            "transaction_id": transaction_id
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to purchase credits: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/credits/{wallet_address}/bonus")
async def claim_welcome_bonus(
    wallet_address: str,
    user_repo=Depends(get_user_repository)
):
    """Claim welcome bonus credits."""
    try:
        user = user_repo.get_user_by_wallet(wallet_address)
        
        if not user:
            # Create user with welcome bonus
            user = user_repo.create_user(
                wallet_address=wallet_address,
                user_type="student",
                subscription_tier="free",
                credits_balance=3,  # Welcome bonus
                welcome_bonus_claimed=True
            )
            
            return {
                "status": "success",
                "credits_granted": 3,
                "new_balance": 3,
                "message": "Welcome bonus claimed!"
            }
        
        # Check if bonus already claimed
        if user.welcome_bonus_claimed:
            raise HTTPException(
                status_code=400,
                detail="Welcome bonus already claimed"
            )
        
        # Grant welcome bonus
        new_balance = user.credits_balance + 3
        user_repo.update_user_stats(
            str(user.id),
            credits_balance=new_balance,
            welcome_bonus_claimed=True
        )
        
        logger.info(f"Welcome bonus claimed: {wallet_address}")
        
        return {
            "status": "success",
            "credits_granted": 3,
            "new_balance": new_balance,
            "message": "Welcome bonus claimed!"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to claim welcome bonus: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Webhook endpoints
@app.post("/api/webhook/dynamic")
async def dynamic_webhook(payload: Dict[str, Any]):
    """Handle Dynamic.xyz webhooks for payment verification."""
    try:
        # Process Dynamic.xyz webhook
        event_type = payload.get("type")
        
        if event_type == "payment.completed":
            # Handle successful payment
            logger.info(f"Payment completed: {payload}")
            
        elif event_type == "user.authenticated":
            # Handle user authentication
            logger.info(f"User authenticated: {payload}")
        
        return {"status": "received"}
        
    except Exception as e:
        logger.error(f"Dynamic webhook processing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/webhook/turnitin")
async def turnitin_webhook(payload: Dict[str, Any]):
    """Handle Turnitin webhooks for plagiarism reports."""
    try:
        # Process Turnitin webhook
        submission_id = payload.get("submission_id")
        status = payload.get("status")
        
        if status == "completed":
            # Handle completed plagiarism check
            logger.info(f"Turnitin check completed for submission: {submission_id}")
        
        return {"status": "received"}
        
    except Exception as e:
        logger.error(f"Turnitin webhook processing failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Import vector storage dependencies
from services.vector_storage import get_vector_storage
from services.embedding_service import get_embedding_service


# Vector search endpoints
@app.post("/api/search/semantic")
async def semantic_search(
    request: Dict[str, Any],
    vector_storage=Depends(get_vector_storage),
    embedding_service=Depends(get_embedding_service)
):
    """Perform semantic search using vector embeddings."""
    try:
        query = request.get("query", "")
        if not query:
            raise HTTPException(status_code=400, detail="Query is required")
        
        # Optional parameters
        academic_field = request.get("academic_field")
        min_credibility = request.get("min_credibility", 0.6)
        limit = min(request.get("limit", 10), 50)  # Max 50 results
        year_range = request.get("year_range")  # [start_year, end_year]
        
        # Generate query embedding
        query_embedding = await embedding_service.embed_query(query, "academic_search")
        
        # Perform semantic search
        results = await vector_storage.semantic_search_documents(
            query_embedding=query_embedding,
            limit=limit,
            academic_field=academic_field,
            min_credibility=min_credibility,
            year_range=tuple(year_range) if year_range and len(year_range) == 2 else None
        )
        
        return {
            "query": query,
            "results": results,
            "total_found": len(results),
            "search_metadata": {
                "academic_field": academic_field,
                "min_credibility": min_credibility,
                "year_range": year_range,
                "timestamp": time.time()
            }
        }
        
    except Exception as e:
        logger.error(f"Semantic search failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/evidence/{conversation_id}")
async def get_conversation_evidence(
    conversation_id: str,
    vector_storage=Depends(get_vector_storage)
):
    """Get all evidence for a conversation with vector similarity support."""
    try:
        evidence_list = await vector_storage.get_conversation_evidence(conversation_id)
        
        return {
            "conversation_id": conversation_id,
            "evidence_count": len(evidence_list),
            "evidence": evidence_list
        }
        
    except Exception as e:
        logger.error(f"Failed to get conversation evidence: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/download/{conversation_id}")
async def list_available_downloads(
    conversation_id: str,
    document_repo=Depends(get_document_repository)
):
    """List all available download formats for a conversation."""
    try:
        documents = document_repo.get_conversation_documents(conversation_id)
        
        if not documents:
            raise HTTPException(
                status_code=404,
                detail=f"No documents found for conversation {conversation_id}"
            )
        
        available_downloads = []
        
        for doc in documents:
            download_info = {
                "document_id": str(doc.id),
                "title": doc.title,
                "document_type": doc.document_type.value if doc.document_type else "unknown",
                "word_count": doc.word_count,
                "quality_score": doc.overall_quality_score,
                "available_formats": []
            }
            
            # Check available formats
            if doc.content_markdown:
                download_info["available_formats"].append({
                    "format": "txt",
                    "url": f"/api/download/{conversation_id}/txt",
                    "description": "Plain text format"
                })
            
            if doc.docx_url or doc.content_markdown:
                download_info["available_formats"].append({
                    "format": "docx", 
                    "url": f"/api/download/{conversation_id}/docx",
                    "description": "Microsoft Word format"
                })
            
            if doc.pdf_url:
                download_info["available_formats"].append({
                    "format": "pdf",
                    "url": f"/api/download/{conversation_id}/pdf", 
                    "description": "PDF format"
                })
            
            if doc.learning_outcomes_coverage:
                download_info["available_formats"].append({
                    "format": "lo_report",
                    "url": f"/api/download/{conversation_id}/lo_report",
                    "description": "Learning outcomes report (JSON)"
                })
            
            available_downloads.append(download_info)
        
        return {
            "conversation_id": conversation_id,
            "documents": available_downloads
        }
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed to list downloads: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/download/{conversation_id}/zip")
async def download_all_formats(
    conversation_id: str,
    document_repo=Depends(get_document_repository)
):
    """Download all available formats as a ZIP file."""
    try:
        import zipfile
        from io import BytesIO
        from docx import Document as DocxDocument
        
        documents = document_repo.get_conversation_documents(conversation_id)
        
        if not documents:
            raise HTTPException(
                status_code=404,
                detail=f"No documents found for conversation {conversation_id}"
            )
        
        # Create ZIP file in memory
        zip_buffer = BytesIO()
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            
            for doc in documents:
                base_name = f"{doc.title.replace(' ', '_')[:50]}"
                
                # Add TXT format
                if doc.content_markdown:
                    zip_file.writestr(f"{base_name}.txt", doc.content_markdown)
                
                # Add DOCX format
                if doc.content_markdown:
                    docx_doc = DocxDocument()
                    for paragraph in doc.content_markdown.split('\n\n'):
                        if paragraph.strip():
                            docx_doc.add_paragraph(paragraph.strip())
                    
                    docx_buffer = BytesIO()
                    docx_doc.save(docx_buffer)
                    zip_file.writestr(f"{base_name}.docx", docx_buffer.getvalue())
                
                # Add learning outcomes report
                if doc.learning_outcomes_coverage:
                    import json
                    lo_json = json.dumps(doc.learning_outcomes_coverage, indent=2)
                    zip_file.writestr(f"{base_name}_learning_outcomes.json", lo_json)
                
                # Add metadata
                metadata = {
                    "title": doc.title,
                    "word_count": doc.word_count,
                    "quality_score": doc.overall_quality_score,
                    "citation_count": doc.citation_count,
                    "academic_field": doc.academic_field,
                    "generated_at": doc.created_at.isoformat() if doc.created_at else None
                }
                zip_file.writestr(f"{base_name}_metadata.json", json.dumps(metadata, indent=2))
        
        zip_buffer.seek(0)
        
        from fastapi.responses import StreamingResponse
        return StreamingResponse(
            BytesIO(zip_buffer.getvalue()),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={conversation_id}_complete.zip"}
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"ZIP download failed: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# Background workflow execution with comprehensive error handling
@with_retry(ErrorCategory.AGENT_FAILURE)
async def execute_writing_workflow(conversation_id: str, initial_state: HandyWriterzState):
    """Execute the writing workflow with production-grade error handling."""
    context = ErrorContext(
        conversation_id=conversation_id,
        user_id=initial_state.user_id,
        additional_data={
            "workflow_type": "academic_writing",
            "user_params": initial_state.user_params
        }
    )
    
    try:
        logger.info(f"🚀 Starting revolutionary workflow for conversation: {conversation_id}")
        
        # Broadcast workflow start with enhanced data
        await redis_client.publish(
            f"sse:{conversation_id}",
            json.dumps({
                "type": "workflow_start",
                "timestamp": time.time(),
                "data": {
                    "conversation_id": conversation_id,
                    "user_id": initial_state.user_id,
                    "estimated_duration": "8-12 minutes",
                    "workflow_version": "2.0.0"
                }
            })
        )
        
        # TODO: Execute the LangGraph workflow with circuit breaker
        # Temporarily disabled until handywriterz_graph is implemented
        # config = {"configurable": {"thread_id": conversation_id}}
        
        workflow_start_time = time.time()
        chunk_count = 0
        
        # Simulate workflow progress for testing
        workflow_steps = [
            ("planning", "Creating document outline"),
            ("researching", "Gathering academic sources"),
            ("filtering", "Validating source credibility"),
            ("writing", "Generating academic content"),
            ("evaluating", "Quality assessment"),
            ("formatting", "Final document formatting")
        ]
        
        for i, (step, description) in enumerate(workflow_steps):
            await asyncio.sleep(2)  # Simulate processing time
            chunk_count += 1
            
            # Enhanced progress broadcasting
            await redis_client.publish(
                f"sse:{conversation_id}",
                json.dumps({
                    "type": "workflow_progress",
                    "timestamp": time.time(),
                    "data": {
                        "current_node": step,
                        "description": description,
                        "progress": (i + 1) / len(workflow_steps) * 100,
                        "chunk_number": chunk_count,
                        "elapsed_time": time.time() - workflow_start_time
                    }
                })
            )
            
            # Log major progress milestones
            logger.info(f"📍 Workflow [{conversation_id}] progressed to: {step}")
        
        # TODO: Uncomment when handywriterz_graph is ready
        # async for chunk in handywriterz_graph.astream(initial_state, config):
        #     chunk_count += 1
        #
        #     # Enhanced progress broadcasting
        #     await redis_client.publish(
        #         f"sse:{conversation_id}",
        #         json.dumps({
        #             "type": "workflow_progress",
        #             "timestamp": time.time(),
        #             "data": {
        #                 **chunk,
        #                 "chunk_number": chunk_count,
        #                 "elapsed_time": time.time() - workflow_start_time
        #             }
        #         })
        #     )
        #
        #     # Log major progress milestones
        #     if "current_node" in chunk:
        #         logger.info(f"📍 Workflow [{conversation_id}] progressed to: {chunk['current_node']}")
        
        workflow_duration = time.time() - workflow_start_time
        
        # Broadcast successful completion
        await redis_client.publish(
            f"sse:{conversation_id}",
            json.dumps({
                "type": "workflow_complete",
                "timestamp": time.time(),
                "data": {
                    "conversation_id": conversation_id,
                    "status": "completed",
                    "duration_seconds": workflow_duration,
                    "chunks_processed": chunk_count,
                    "completion_message": "✅ Academic document generated successfully! (Demo mode)"
                }
            })
        )
        
        logger.info(f"✅ Workflow completed successfully for {conversation_id} in {workflow_duration:.2f}s")
        
    except Exception as e:
        workflow_duration = time.time() - workflow_start_time if 'workflow_start_time' in locals() else 0
        
        # Handle error through error handler
        error_data = await error_handler.handle_error(
            e, context, ErrorCategory.AGENT_FAILURE, ErrorSeverity.HIGH
        )
        
        logger.error(f"❌ Workflow execution failed for {conversation_id}: {e}")
        
        # Determine if error is recoverable
        recovery_strategy = error_data.get("recovery_strategy", {})
        is_recoverable = recovery_strategy.get("retry_recommended", False)
        
        # Broadcast workflow failure with recovery information
        await redis_client.publish(
            f"sse:{conversation_id}",
            json.dumps({
                "type": "workflow_failed",
                "timestamp": time.time(),
                "data": {
                    "conversation_id": conversation_id,
                    "error": str(e),
                    "error_id": error_data.get("error_id"),
                    "error_type": type(e).__name__,
                    "duration_seconds": workflow_duration,
                    "recoverable": is_recoverable,
                    "recovery_strategy": recovery_strategy,
                    "support_message": "Our team has been notified. Please try again or contact support."
                }
            })
        )
        
        # Re-raise if not recoverable
        if not is_recoverable:
            raise


# Configuration endpoint for frontend compatibility
@app.get("/api/config")
@with_error_handling(ErrorCategory.SYSTEM, ErrorSeverity.LOW)
async def get_app_config(request: Request):
    """
    Get application configuration for frontend.
    Compatible with OpenWebUI frontend expectations.
    """
    
    # For now, return a simplified config that satisfies the frontend
    # This can be expanded as needed
    return {
        "status": True,
        "name": "HandyWriterz",
        "version": "2.0.0",
        "default_locale": "en-US",
        "oauth": {
            "providers": {}
        },
        "features": {
            "auth": True,
            "auth_trusted_header": False,
            "enable_ldap": False,
            "enable_api_key": True,
            "enable_signup": True,
            "enable_login_form": True,
            "enable_websocket": True,
            "enable_direct_connections": False,
            "enable_channels": False,
            "enable_notes": True,
            "enable_web_search": True,
            "enable_code_execution": False,
            "enable_code_interpreter": False,
            "enable_image_generation": False,
            "enable_autocomplete_generation": True,
            "enable_community_sharing": False,
            "enable_message_rating": True,
            "enable_user_webhooks": False,
            "enable_admin_export": False,
            "enable_admin_chat_access": True,
            "enable_google_drive_integration": False,
            "enable_onedrive_integration": False,
        },
        "default_models": [],
        "default_prompt_suggestions": [
            {
                "title": "Academic Essay",
                "content": "Help me write a 2000-word academic essay on sustainable healthcare practices with Harvard referencing."
            },
            {
                "title": "Research Report", 
                "content": "Create a comprehensive research report analyzing the impact of AI on modern education systems."
            },
            {
                "title": "Literature Review",
                "content": "Write a literature review examining recent developments in renewable energy technologies."
            }
        ],
        "user_count": 1,
        "code": {
            "engine": "none"
        },
        "audio": {
            "tts": {
                "engine": "none",
                "voice": "default",
                "split_on": "punctuation"
            },
            "stt": {
                "engine": "none"
            }
        },
        "file": {
            "max_size": 50,  # MB
            "max_count": 10,
            "image_compression": {
                "width": 1920,
                "height": 1080
            }
        },
        "permissions": {
            "workspace": {
                "models": True,
                "knowledge": True,
                "prompts": True,
                "tools": True,
                "functions": True
            },
            "chat": {
                "file_upload": True,
                "delete": True,
                "edit": True,
                "temporary": True
            }
        }
    }


# SPA Catch-all route to serve the frontend
@app.get("/")
async def serve_frontend():
    """Serve the SvelteKit build index.html file."""
    build_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "build")
    index_html = os.path.join(build_dir, "index.html")
    
    if os.path.exists(index_html):
        return FileResponse(index_html, media_type="text/html")
    else:
        # Fallback to development mode - serve from src
        frontend_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(__file__))), "src")
        app_html = os.path.join(frontend_dir, "app.html")
        
        if os.path.exists(app_html):
            return FileResponse(app_html, media_type="text/html")
        else:
            return JSONResponse(
                status_code=503,
                content={
                    "error": "Frontend not available",
                    "message": "Please build the frontend with 'npm run build' or run in development mode"
                }
            )

@app.get("/chat")
async def serve_chat():
    """Serve the chat page."""
    return await serve_frontend()

@app.get("/c/{chat_id}")
async def serve_chat_with_id(chat_id: str):
    """Serve specific chat page."""
    return await serve_frontend()

@app.get("/{path:path}")
async def serve_spa(path: str):
    """Catch-all route for SPA routing."""
    # Skip API routes
    if path.startswith("api/"):
        return JSONResponse(status_code=404, content={"error": "API endpoint not found"})
    
    # Skip static files
    if path.startswith("static/") or path.startswith("_app/") or path.startswith("pyodide/"):
        return JSONResponse(status_code=404, content={"error": "Static file not found"})
    
    # For all other routes, serve the SPA
    return await serve_frontend()


if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=8000,
        reload=True,
        log_level="info"
    )